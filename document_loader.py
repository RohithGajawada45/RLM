"""
Document loader — extracts text from PDFs and DOCX files.

PDF extraction priority:
    1. pdftotext -layout (preserves tables; needed for financial docs)
    2. pypdf (pure-Python fallback)

DOCX extraction uses python-docx.
"""

import os
import subprocess
import tempfile
from pathlib import Path
from typing import Optional


def load_document(path: str) -> str:
    """Dispatch on file extension. Returns clean text as a single string."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    ext = p.suffix.lower()
    if ext == ".pdf":
        return _clean_text(_load_pdf(path))
    if ext == ".docx":
        return _clean_text(_load_docx(path))
    if ext in (".txt", ".md"):
        return _clean_text(
            Path(path).read_text(encoding="utf-8", errors="replace")
        )
    raise ValueError(
        f"Unsupported file type: {ext}. Use .pdf, .docx, .txt, or .md"
    )


# ─── PDF extraction ───────────────────────────────────────────────

def _load_pdf(path: str) -> str:
    """Try pdftotext first, fall back to pypdf."""
    text = _try_pdftotext(path)
    if text:
        return text
    return _load_pdf_pypdf(path)

def _try_pdftotext(path: str) -> Optional[str]:
    """Use pdftotext CLI if available. Preserves table layout for financial docs."""
    try:
        with tempfile.NamedTemporaryFile(
            mode="r", suffix=".txt",
            delete=False, encoding="utf-8"
        ) as f:
            tmp_path = f.name

        result = subprocess.run(
            ["pdftotext", "-layout", path, tmp_path],
            capture_output=True, timeout=120,
        )
        if result.returncode != 0:
            return None

        with open(tmp_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

        os.unlink(tmp_path)
        return text

    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        return None


def _load_pdf_pypdf(path: str) -> str:
    """Pure-Python PDF fallback."""
    import pypdf

    reader = pypdf.PdfReader(path)
    pages = []

    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append(f"[Page {i+1} extraction failed]")

        pages.append("\f")  # form-feed = page break marker

    return "".join(pages)


# ─── DOCX extraction ───────────────────────────────────────────────

def _load_docx(path: str) -> str:
    """Extract paragraphs + table cells from a .docx file using python-docx."""
    try:
        import docx  # python-docx
    except ImportError:
                raise ImportError(
            "python-docx is required for .docx files. Install with: pip install python-docx"
        )

    doc = docx.Document(path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Pull tables too — DOCX often hides data there
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n".join(parts)


# ─── Text cleanup ─────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """Normalize whitespace; convert form-feeds into explicit page break markers."""
    if not text:
        return ""

    text = text.replace("\f", "\n\n--- PAGE BREAK ---\n\n")

    # Collapse 3+ blank lines into 2
    out_lines, blanks = [], 0

    for line in text.splitlines():
        stripped = line.rstrip()

        if not stripped:
            blanks += 1
            if blanks <= 2:
                out_lines.append("")
        else:
            blanks = 0
            out_lines.append(stripped)

    return "\n".join(out_lines).strip()


def get_document_stats(text: str) -> dict:
    return {
        "total_chars": len(text),
        "total_words": len(text.split()),
        "total_lines": text.count("\n") + 1,
        "estimated_pages": max(1, len(text) // 3000),
    }


# Backwards compatibility
def load_pdf_text(path: str) -> str:
    return load_document(path)